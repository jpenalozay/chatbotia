"""
Servicio RAG (Retrieval-Augmented Generation)
Maneja el procesamiento de documentos, embeddings y consultas RAG
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
from datetime import datetime

# Document processing
import fitz  # PyMuPDF
from docx import Document
import pandas as pd

# RAG components
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document as LangchainDocument

# Database
from sqlalchemy.orm import Session
from app.models.current import SystemUser
import chromadb
from chromadb.config import Settings

logger = logging.getLogger(__name__)


class RAGService:
    """Servicio principal para RAG"""
    
    def __init__(self, persist_directory: str = "./chroma_db"):
        """
        Inicializa el servicio RAG
        
        Args:
            persist_directory: Directorio para persistir la base de datos vectorial
        """
        self.persist_directory = persist_directory
        os.makedirs(persist_directory, exist_ok=True)
        
        # Inicializar ChromaDB con telemetría desactivada para evitar errores de captura
        self.chroma_client = chromadb.PersistentClient(
            path=persist_directory,
            settings=Settings(anonymized_telemetry=False)
        )
        
        # Inicializar modelo de embeddings
        logger.info("Inicializando modelo de embeddings...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        
        # Inicializar vector store
        logger.info("Inicializando Chroma vector store...")
        self.vector_store = Chroma(
            client=self.chroma_client,
            embedding_function=self.embeddings,
            collection_name="user_documents"
        )
        
        # Text splitter para chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=512,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        logger.info("RAG Service inicializado correctamente")
    
    def process_pdf(self, file_path: str) -> str:
        """
        Extrae texto de un archivo PDF
        
        Args:
            file_path: Ruta al archivo PDF
            
        Returns:
            Texto extraído del PDF
        """
        try:
            text = ""
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            logger.error(f"Error procesando PDF {file_path}: {e}")
            raise
    
    def process_docx(self, file_path: str) -> str:
        """
        Extrae texto de un archivo Word (.docx)
        
        Args:
            file_path: Ruta al archivo DOCX
            
        Returns:
            Texto extraído del documento
        """
        try:
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error procesando DOCX {file_path}: {e}")
            raise
    
    def process_excel(self, file_path: str) -> str:
        """
        Extrae texto de un archivo Excel (.xlsx)
        
        Args:
            file_path: Ruta al archivo Excel
            
        Returns:
            Texto extraído del Excel (formato CSV-like)
        """
        try:
            df = pd.read_excel(file_path, sheet_name=None)  # Lee todas las hojas
            text = ""
            for sheet_name, sheet_df in df.items():
                text += f"\n\n=== Hoja: {sheet_name} ===\n"
                text += sheet_df.to_string(index=False)
            return text
        except Exception as e:
            logger.error(f"Error procesando Excel {file_path}: {e}")
            raise
    
    def process_txt(self, file_path: str) -> str:
        """
        Lee un archivo de texto plano
        
        Args:
            file_path: Ruta al archivo TXT
            
        Returns:
            Contenido del archivo
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error procesando TXT {file_path}: {e}")
            raise
    
    def process_document(self, file_path: str, file_type: str) -> str:
        """
        Procesa un documento según su tipo
        
        Args:
            file_path: Ruta al archivo
            file_type: Tipo de archivo (pdf, docx, xlsx, txt, md)
            
        Returns:
            Texto extraído del documento
        """
        processors = {
            'pdf': self.process_pdf,
            'docx': self.process_docx,
            'xlsx': self.process_excel,
            'txt': self.process_txt,
            'md': self.process_txt,
        }
        
        processor = processors.get(file_type.lower())
        if not processor:
            raise ValueError(f"Tipo de archivo no soportado: {file_type}")
        
        return processor(file_path)
    
    def chunk_text(self, text: str, chunk_size: int = 512, chunk_overlap: int = 50) -> List[str]:
        """
        Divide el texto en chunks
        
        Args:
            text: Texto a dividir
            chunk_size: Tamaño de cada chunk
            chunk_overlap: Overlap entre chunks
            
        Returns:
            Lista de chunks de texto
        """
        # Actualizar configuración del splitter
        self.text_splitter.chunk_size = chunk_size
        self.text_splitter.chunk_overlap = chunk_overlap
        
        chunks = self.text_splitter.split_text(text)
        return chunks
    
    def add_document_to_vectorstore(
        self,
        chunks: List[str],
        system_user_id: int,
        company_id: int,
        document_id: int,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[str]:
        """
        Agrega chunks de un documento al vector store
        
        Args:
            chunks: Lista de chunks de texto
            system_user_id: ID del usuario (SystemUser)
            company_id: ID de la empresa
            document_id: ID del documento
            filename: Nombre del archivo
            metadata: Metadata adicional
            
        Returns:
            Lista de IDs de los chunks en el vector store
        """
        try:
            # Preparar documentos para LangChain
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    "user_id": system_user_id,
                    "company_id": company_id,
                    "document_id": document_id,
                    "filename": filename,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "timestamp": datetime.now().isoformat()
                }
                if metadata:
                    doc_metadata.update(metadata)
                
                documents.append(
                    LangchainDocument(
                        page_content=chunk,
                        metadata=doc_metadata
                    )
                )
            
            # Agregar al vector store
            ids = self.vector_store.add_documents(documents)
            
            # ChromaDB 0.4.x persiste automáticamente, no necesita persist() manual
            
            logger.info(f"Agregados {len(chunks)} chunks del documento {document_id} al vector store")
            return ids
            
        except Exception as e:
            logger.error(f"Error agregando documento al vector store: {e}")
            raise
    
    def search_similar_chunks(
        self,
        query: str,
        user_id: Optional[int] = None,
        company_id: Optional[int] = None,
        top_k: int = 5,
        filter_metadata: Optional[Dict[str, Any]] = None,
        **kwargs
    ) -> List[Dict[str, Any]]:
        """
        Busca chunks similares a una query

        Args:
            query: Query de búsqueda
            user_id: ID del usuario (opcional)
            company_id: ID de la empresa (opcional)
            top_k: Número de resultados a retornar
            filter_metadata: Filtros adicionales de metadata
            
        Returns:
            Lista de chunks relevantes con metadata
        """
        if kwargs:
            logger.warning(f"⚠️ search_similar_chunks recibió argumentos inesperados: {kwargs}")
            if 'system_user_id' in kwargs and not user_id:
                user_id = kwargs['system_user_id']
        try:
            # Preparar filtros dinámicamente según ChromaDB ($and si hay múltiples)
            filter_list = []
            if user_id:
                filter_list.append({"user_id": user_id})
            if company_id:
                filter_list.append({"company_id": company_id})
            
            if filter_metadata:
                for k, v in filter_metadata.items():
                    filter_list.append({k: v})

            final_filter = None
            if len(filter_list) == 1:
                final_filter = filter_list[0]
            elif len(filter_list) > 1:
                final_filter = {"$and": filter_list}
        
            # Buscar documentos similares
            results = self.vector_store.similarity_search_with_score(
                query=query,
                k=top_k,
                filter=final_filter
            )
            
            # Formatear resultados
            formatted_results = []
            for doc, score in results:
                formatted_results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "similarity_score": float(score)
                })
            
            logger.info(f"Encontrados {len(formatted_results)} chunks (user_id={user_id}, company_id={company_id})")
            return formatted_results
            
        except Exception as e:
            logger.error(f"Error buscando chunks similares: {e}")
            raise
    
    def delete_document_from_vectorstore(self, document_id: int) -> bool:
        """
        Elimina todos los chunks de un documento del vector store
        
        Args:
            document_id: ID del documento a eliminar
            
        Returns:
            True si se eliminó correctamente
        """
        try:
            # Eliminar del vector store usando el document_id en la metadata
            self.vector_store.delete(where={"document_id": document_id})
            logger.info(f"✅ Documento {document_id} eliminado del vector store")
            return True
        except Exception as e:
            logger.error(f"Error eliminando documento del vector store: {e}")
            return False
    
    def get_company_document_count(self, company_id: int) -> int:
        """
        Obtiene el número de documentos de una empresa en el vector store
        
        Args:
            company_id: ID de la empresa
            
        Returns:
            Número de documentos
        """
        try:
            results = self.vector_store.similarity_search(
                query="",
                k=1000,
                filter={"company_id": company_id}
            )
            return len(results)
        except Exception as e:
            logger.error(f"Error contando documentos de empresa: {e}")
            return 0


# Instancia global del servicio RAG
rag_service = RAGService()
